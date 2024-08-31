import streamlit as st
from docx import Document
from fpdf import FPDF
import pypandoc
import re
import os

def main():
    st.title("Premier Instruction Manual Viewer and Editor")

    st.sidebar.title("Navigation")
    section = st.sidebar.selectbox("Select Section", [
        "Introduction", 
        "Upload and View Manual", 
        "Edit Manual", 
        "Search Manual", 
        "Export Manual"
    ])

    if section == "Introduction":
        show_introduction()
    elif section == "Upload and View Manual":
        upload_and_view_manual()
    elif section == "Edit Manual":
        edit_manual()
    elif section == "Search Manual":
        search_manual()
    elif section == "Export Manual":
        export_manual()

def show_introduction():
    st.header("Introduction")
    st.write("""
    Welcome to the Premier Instruction Manual Viewer and Editor. 
    You can upload manuals in various formats, edit them, navigate through sections, search for specific terms, and export the content to different formats such as PDF or DOCX.
    """)

def upload_and_view_manual():
    st.header("Upload and View Manual")
    uploaded_file = st.file_uploader("Choose a file", type=["docx", "txt", "pdf", "rtf"])

    if uploaded_file is not None:
        file_extension = os.path.splitext(uploaded_file.name)[1].lower()

        if file_extension == '.txt':
            content = uploaded_file.read().decode("utf-8")
        elif file_extension == '.docx':
            content = read_docx(uploaded_file)
        elif file_extension == '.pdf':
            content = read_pdf(uploaded_file)
        elif file_extension == '.rtf':
            content = read_rtf(uploaded_file)
        else:
            st.error("Unsupported file format.")
            return

        st.session_state['manual_content'] = content
        sections = parse_sections(content)
        if sections:
            selected_section = st.sidebar.radio("Jump to Section", list(sections.keys()))
            st.subheader(selected_section)
            st.write(sections[selected_section])

def read_docx(uploaded_file):
    """Reads a .docx file and returns its content as plain text."""
    doc = Document(uploaded_file)
    full_text = []
    for paragraph in doc.paragraphs:
        full_text.append(paragraph.text)
    return '\n'.join(full_text)

def read_pdf(uploaded_file):
    """Reads a .pdf file and returns its content as plain text."""
    reader = PdfReader(uploaded_file)
    full_text = []
    for page in reader.pages:
        full_text.append(page.extract_text())
    return '\n'.join(full_text)

def read_rtf(uploaded_file):
    """Reads a .rtf file and returns its content as plain text."""
    output = pypandoc.convert_file(uploaded_file.name, 'plain', format='rtf')
    return output

def parse_sections(content):
    """Parses the content into sections based on headings."""
    sections = {}
    lines = content.splitlines()
    current_section = None
    section_text = []

    for line in lines:
        if is_heading(line):
            if current_section:
                sections[current_section] = '\n'.join(section_text)
                section_text = []
            current_section = line.strip()
        else:
            section_text.append(line)

    if current_section:
        sections[current_section] = '\n'.join(section_text)
    
    return sections

def is_heading(line):
    """Determines if a line is a heading."""
    return re.match(r'(Chapter|Section)\s+\d+', line, re.IGNORECASE) or len(line.split()) < 5

def edit_manual():
    st.header("Edit Manual")
    
    if 'manual_content' not in st.session_state:
        st.error("Please upload and view a manual first.")
        return
    
    edited_content = st.text_area("Edit the manual content:", st.session_state['manual_content'], height=300)
    if st.button("Save Changes"):
        st.session_state['manual_content'] = edited_content
        st.success("Changes saved.")

def search_manual():
    st.header("Search Manual")
    
    if 'manual_content' not in st.session_state:
        st.error("Please upload and view a manual first.")
        return
    
    search_term = st.text_input("Search for a term or phrase:")
    if search_term:
        search_results = search_sections(st.session_state['manual_content'], search_term)
        if search_results:
            for section, content in search_results.items():
                st.subheader(section)
                st.write(content, unsafe_allow_html=True)
        else:
            st.write("No results found.")

def search_sections(content, search_term):
    """Searches for a term within sections and returns a dictionary of matched sections with highlighted terms."""
    sections = parse_sections(content)
    search_results = {}
    for section, content in sections.items():
        if search_term.lower() in content.lower():
            highlighted_content = highlight_term(content, search_term)
            search_results[section] = highlighted_content
    return search_results

def highlight_term(content, term):
    """Highlights all occurrences of the search term in the content."""
    highlighted = re.sub(f"({re.escape(term)})", r'<mark>\1</mark>', content, flags=re.IGNORECASE)
    return highlighted

def export_manual():
    st.header("Export Manual")
    
    if 'manual_content' not in st.session_state:
        st.error("Please upload and edit a manual first.")
        return
    
    export_format = st.radio("Select export format:", ("DOCX", "PDF"))
    
    if export_format == "DOCX":
        export_as_docx(st.session_state['manual_content'])
    elif export_format == "PDF":
        export_as_pdf(st.session_state['manual_content'])

def export_as_docx(content):
    """Exports the content as a DOCX file."""
    doc = Document()
    for line in content.splitlines():
        doc.add_paragraph(line)
    
    doc_file = "exported_manual.docx"
    doc.save(doc_file)
    
    with open(doc_file, "rb") as f:
        st.download_button("Download DOCX", f, file_name=doc_file, mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document")

def export_as_pdf(content):
    """Exports the content as a PDF file."""
    pdf = FPDF()
    pdf.add_page()
    pdf.set_font("Arial", size=12)
    
    for line in content.splitlines():
        pdf.multi_cell(0, 10, line)
    
    pdf_file = "exported_manual.pdf"
    pdf.output(pdf_file)
    
    with open(pdf_file, "rb") as f:
        st.download_button("Download PDF", f, file_name=pdf_file, mime="application/pdf")

if __name__ == "__main__":
    main()
